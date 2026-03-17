import json
import os
import re
import shutil
from docx import Document
from openai import OpenAI


def _get_client():
    api_key = os.environ.get("OPENROUTER_API_KEY")
    if not api_key:
        raise ValueError("OPENROUTER_API_KEY environment variable is not set.")
    return OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)


def _para_text(para):
    return "".join(run.text for run in para.runs)


def _set_para_text(para, new_text: str):
    """Replace paragraph text while preserving all run-level formatting.
    Puts new text in first run, blanks out the rest."""
    if not para.runs:
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _is_modifiable_text(text: str) -> bool:
    """Identify text safe to reword: bullet points and descriptive lines.
    Skip: empty lines, short headers, contact info, ALL CAPS headers."""
    text = text.strip()
    if not text:
        return False
    # Skip very short lines (likely headers/names)
    if len(text.split()) <= 3:
        return False
    # Skip lines that look like contact info (email, phone, urls)
    # FIX: use alternation (|) not character class ([])
    if re.search(r"@|linkedin\.com|github\.com|https?://|\+?\d[\d\s\-\(\)]{7,}", text, re.I):
        return False
    # Skip lines that are ALL CAPS (section headers)
    if text.isupper():
        return False
    return True


def extract_modifiable(doc: Document) -> dict[int, str]:
    """Return {paragraph_index: text} for paragraphs that can be reworded."""
    result = {}
    for i, para in enumerate(doc.paragraphs):
        text = _para_text(para).strip()
        if _is_modifiable_text(text):
            result[i] = text
    return result


def tailor_with_claude(paragraphs: dict[int, str], jd: str) -> dict[int, str]:
    """Send resume paragraphs + JD to Claude, get back reworded versions."""
    para_list = "\n".join(
        f'[{idx}] {text}' for idx, text in paragraphs.items()
    )

    prompt = f"""You are a professional resume writer. Your job is to tailor resume bullet points and descriptions to better match a job description — WITHOUT changing the format, structure, or length significantly.

RULES:
- Rewrite each line to naturally incorporate relevant keywords/skills from the JD
- Keep roughly the same length and sentence structure
- Do NOT add new bullet points or remove existing ones
- Do NOT change section headers, names, dates, or contact info (they won't be sent)
- Keep quantifiable achievements (numbers, percentages) intact
- Sound natural and professional, not keyword-stuffed
- If a line is already a good match, keep it nearly the same

JOB DESCRIPTION:
{jd}

RESUME LINES TO TAILOR (format: [index] text):
{para_list}

Return a JSON object mapping each index to its tailored version. Example:
{{"0": "tailored line here", "5": "another tailored line"}}

Return ONLY valid JSON, no explanation."""

    message = _get_client().chat.completions.create(
        model="google/gemini-2.0-flash-001",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )

    raw = message.choices[0].message.content.strip()
    # Strip markdown code fences if present
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)

    data = json.loads(raw)
    return {int(k): v for k, v in data.items()}


def apply_tailoring(doc: Document, changes: dict[int, str]) -> Document:
    """Apply Claude's rewrites back to the document."""
    for idx, new_text in changes.items():
        if idx < len(doc.paragraphs):
            _set_para_text(doc.paragraphs[idx], new_text)
    return doc


def tailor_resume_docx(input_path: str, output_path: str, jd: str):
    doc = Document(input_path)
    modifiable = extract_modifiable(doc)
    if not modifiable:
        doc.save(output_path)
        return
    changes = tailor_with_claude(modifiable, jd)
    apply_tailoring(doc, changes)
    doc.save(output_path)


def tailor_resume_pdf(input_path: str, output_path: str, jd: str):
    """Tailor a PDF resume while preserving the original format.
    Uses PyMuPDF to do in-place text replacement on the PDF."""
    import fitz

    doc = fitz.open(input_path)

    # Step 1: Extract text blocks with position and font metadata
    blocks_info = []
    for page_idx in range(len(doc)):
        page = doc[page_idx]
        page_dict = page.get_text("dict")
        for block in page_dict["blocks"]:
            if block["type"] != 0:  # skip image blocks
                continue
            full_text = ""
            first_span = None
            for line in block["lines"]:
                for span in line["spans"]:
                    if first_span is None:
                        first_span = span
                    full_text += span["text"]
                full_text += " "
            full_text = full_text.strip()

            if not full_text or first_span is None:
                continue

            blocks_info.append({
                "page_idx": page_idx,
                "text": full_text,
                "bbox": block["bbox"],
                "fontsize": first_span["size"],
                "color": first_span["color"],
            })

    # Step 2: Identify modifiable blocks
    modifiable = {}
    for i, bi in enumerate(blocks_info):
        if _is_modifiable_text(bi["text"]):
            modifiable[i] = bi["text"]

    if not modifiable:
        shutil.copy(input_path, output_path)
        doc.close()
        return

    # Step 3: Get tailored text from AI
    changes = tailor_with_claude(modifiable, jd)

    # Step 4: Replace text in PDF
    # First pass: add redaction annotations to clear old text
    for idx in changes:
        bi = blocks_info[idx]
        page = doc[bi["page_idx"]]
        rect = fitz.Rect(bi["bbox"])
        page.add_redact_annot(rect, fill=(1, 1, 1))  # white fill

    # Apply all redactions
    for page in doc:
        page.apply_redactions()

    # Second pass: insert new text at same positions
    for idx, new_text in changes.items():
        bi = blocks_info[idx]
        page = doc[bi["page_idx"]]
        rect = fitz.Rect(bi["bbox"])
        fontsize = bi["fontsize"]
        color_int = bi["color"]

        # Convert color int to RGB tuple
        r = ((color_int >> 16) & 0xFF) / 255.0
        g = ((color_int >> 8) & 0xFF) / 255.0
        b = (color_int & 0xFF) / 255.0

        tw = fitz.TextWriter(page.rect)
        font = fitz.Font("helv")
        tw.fill_textbox(rect, new_text, fontsize=fontsize, font=font, align=0)
        tw.write_text(page, color=(r, g, b))

    doc.save(output_path)
    doc.close()


def tailor_resume(input_path: str, output_path: str, jd: str):
    if input_path.lower().endswith(".pdf"):
        tailor_resume_pdf(input_path, output_path, jd)
    else:
        tailor_resume_docx(input_path, output_path, jd)
